# -*- coding: utf-8 -*-
"""Generate 开发说明书.docx using python-docx"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

for level, (size, color) in enumerate([(22, '1A3A2A'), (16, '1A3A2A'), (14, '333333')], 1):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'SimHei'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor.from_string(color)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
    hs.paragraph_format.space_after = Pt(10 if level == 1 else 8)

def add_para(text, style_name='Normal', bold=False, align=None, indent_first=True, font_name=None, font_size=None, color=None, space_after=None, space_before=None):
    para = doc.add_paragraph(style=style_name)
    if align:
        para.alignment = align
    if not indent_first and style_name == 'Normal':
        para.paragraph_format.first_line_indent = Cm(0)
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        para.paragraph_format.space_before = Pt(space_before)
    run = para.add_run(text)
    if bold:
        run.bold = True
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return para

def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_font(cell, bold=False, font_name='SimSun', font_size=10, color=None):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if bold:
                run.bold = True
            if color:
                run.font.color.rgb = RGBColor.from_string(color)

def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1A3A2A')
        set_cell_font(cell, bold=True, font_name='SimHei', font_size=10, color='FFFFFF')

    # Data rows
    for r, row_data in enumerate(rows):
        for i, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[i]
            cell.text = str(val)
            set_cell_font(cell, font_size=10)
            if r % 2 == 0:
                set_cell_shading(cell, 'F5F5F5')

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table

def add_code_block(lines):
    """Add a code-style block using monospace font."""
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ══════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para('智学', align=WD_ALIGN_PARAGRAPH.CENTER, font_name='SimHei', font_size=36, bold=True, color='1A3A2A', space_after=4)
add_para('ZhiXue', align=WD_ALIGN_PARAGRAPH.CENTER, font_name='Arial', font_size=18, color='666666', space_after=12)

# Decorative line
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_after = Pt(20)
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="1A3A2A"/></w:pBdr>')
para._p.get_or_add_pPr().append(pBdr)

add_para('AI 个性化学习系统', align=WD_ALIGN_PARAGRAPH.CENTER, font_name='SimHei', font_size=20, color='1A3A2A', space_after=30)

for _ in range(3):
    doc.add_paragraph()

add_para('软件开发说明书', align=WD_ALIGN_PARAGRAPH.CENTER, font_name='SimHei', font_size=28, bold=True, color='333333', space_after=30)

for _ in range(4):
    doc.add_paragraph()

cover_info = [
    ('项目名称', '智学（ZhiXue）— 基于多智能体的 AI 个性化学习系统'),
    ('参赛赛道', '第十五届中国大学生软件杯 A3 赛题'),
    ('文档版本', 'v2.0'),
    ('编制日期', '2026 年 6 月'),
]
for label, value in cover_info:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(4)
    run1 = para.add_run(f'{label}：')
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.name = 'SimHei'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run2 = para.add_run(value)
    run2.font.size = Pt(12)
    run2.font.name = 'SimSun'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

doc.add_page_break()

# ══════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════
add_para('目  录', align=WD_ALIGN_PARAGRAPH.CENTER, font_name='SimHei', font_size=18, bold=True, color='1A3A2A', space_after=20)

# Add a TOC field
para = doc.add_paragraph()
run = para.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar1)
run2 = para.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
run2._r.append(instrText)
run3 = para.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run3._r.append(fldChar2)
run4 = para.add_run('（请在 Word 中右键点击此处，选择"更新域"以生成目录）')
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run5 = para.add_run()
fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run5._r.append(fldChar3)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════════════════════
add_heading('一、项目概述', level=1)

add_heading('1.1 项目背景', level=2)
add_para('在数字化与智能化深度融合的时代，高等教育的个性化变革成为核心发展方向。不同学生在知识基础、学习能力、兴趣方向上的显著差异，使得标准化教学难以满足个性化学习需求。本项目旨在借助大模型技术体系（通用大模型、多模态生成、AI 辅助编程工具等），构建高等教育个性化学习资源体系，开发智能学习智能体系统，实现“因材施教”的数字化落地。')

add_heading('1.2 项目目标', level=2)
add_para('以“人工智能导论”课程为切入点，构建多智能体系统，实现：')
goals = [
    ('1. 对话式学习画像自主构建', '：通过自然语言对话自动提取≥6 维度的学生画像，支持动态更新'),
    ('2. 多智能体协同资源生成', '：多个 AI Agent 协作生成文档、思维导图、练习题、视频脚本、代码示例等 5 类学习资源'),
    ('3. 个性化学习路径规划', '：基于知识图谱和学生画像，自动生成拓扑排序的个性化学习路径'),
    ('4. 智能辅导', '：基于 RAG 检索增强的多模态答疑，支持来源引用和难度自适应'),
    ('5. 学习效果评估', '：五维雷达图 + AI 评估报告 + 趋势分析'),
]
for bold_part, normal_part in goals:
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    r1 = para.add_run(bold_part)
    r1.bold = True
    r2 = para.add_run(normal_part)

add_heading('1.3 系统特色', level=2)
add_table(
    ['特色', '说明'],
    [
        ['多智能体协作', '4 个 AI Agent 协同工作（RAG 检索→需求分析→专业生成→安全审查）'],
        ['六维动态画像', '知识基础、认知风格、易错点、学习目标、可用时间、兴趣方向'],
        ['知识图谱驱动', '10 章 DAG 拓扑排序 + 画像权重个性化排序'],
        ['SSE 流式输出', '所有 AI 生成接口均支持 Server-Sent Events 流式推送'],
        ['多模态支持', '文本 + LaTeX 数学公式 + Mermaid 图表 + 代码高亮 + 数字人 + 语音'],
        ['双层安全审查', '正则快筛 + LLM-as-Judge 深度审查'],
    ],
    [3.5, 10.5]
)

add_heading('1.4 代码规模', level=2)
add_table(
    ['模块', '文件数', '代码行数'],
    [
        ['后端 Python', '~50', '~10,746'],
        ['前端 TypeScript/TSX', '~55', '~13,183'],
        ['知识库（Markdown）', '12', '~3,959'],
        ['多智能体定义', '13', '~1,127'],
        ['合计', '~130', '~29,015'],
    ],
    [5, 4, 5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 2
# ══════════════════════════════════════════════════
add_heading('二、需求分析', level=1)

add_heading('2.1 用户痛点分析', level=2)
add_table(
    ['痛点', '具体描述', '解决方案'],
    [
        ['学习资源繁杂无序', '网络上 AI 学习资料海量但质量参差不齐', '多智能体协同生成高质量、结构化学习资源'],
        ['缺乏个性化路径', '所有学生学一样的内容，无法匹配个人水平', '六维画像 + 知识图谱权重 + 个性化路径规划'],
        ['薄弱环节难定位', '学生不清楚自己哪些知识点未掌握', '做题诊断 + 画像自动更新易错点'],
        ['疑问无法即时解答', '课后遇到问题无人可问', 'AI 智能辅导（RAG + 流式答疑 + 来源引用）'],
        ['学习动力不足', '缺乏正反馈和成就激励', '每日打卡 + 成就徽章 + 学习趋势可视化'],
    ],
    [3, 5.5, 5.5]
)

add_heading('2.2 功能需求矩阵', level=2)
add_table(
    ['编号', '功能模块', '优先级', '对应赛题要求'],
    [
        ['FR-01', '对话式学习画像构建', 'P0', '≥6 维度，动态更新'],
        ['FR-02', '多智能体协同资源生成', 'P0', '≥5 种类型，不同角色协作'],
        ['FR-03', '个性化学习路径规划', 'P0', '路径规划与资源推送'],
        ['FR-04', '智能辅导答疑', 'P1', '加分项'],
        ['FR-05', '学习效果评估', 'P1', '加分项'],
        ['FR-06', '在线练习与自动批改', 'P1', '资源类型丰富度'],
        ['FR-07', '错题本与复习', 'P1', '学习效果追踪'],
        ['FR-08', '知识库管理', 'P1', 'RAG 检索增强'],
        ['FR-09', 'AI 课堂模拟', 'P2', '智能教学场景'],
        ['FR-10', 'PPT 生成', 'P2', '多模态资源'],
        ['FR-11', '数字人交互', 'P2', '多模态交互'],
        ['FR-12', '语音交互', 'P2', '多模态交互'],
    ],
    [1.5, 4, 2, 6.5]
)

add_heading('2.3 非功能性需求', level=2)
add_table(
    ['需求类型', '具体指标'],
    [
        ['响应效率', 'SSE 流式输出首 token 延迟 < 2s，生成进度实时可见'],
        ['内容安全', '10 类正则 + LLM-as-Judge 双层审查，阻断违规内容'],
        ['渲染能力', 'Markdown + LaTeX 数学公式 + Mermaid 图表 + 代码高亮'],
        ['并发能力', '支持多用户同时使用，SQLite 异步读写'],
        ['可用性', '响应式布局，支持桌面端和移动端'],
    ],
    [3.5, 10.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 3
# ══════════════════════════════════════════════════
add_heading('三、系统总体设计', level=1)

add_heading('3.1 系统架构图', level=2)
add_para('系统采用前后端分离架构，前端通过 REST API 和 SSE 流与后端通信，后端通过多智能体编排引擎调用 LLM 服务和向量检索服务。架构层次如下：')

arch = [
    '┌─────────────────────────────────────────────────────────────────┐',
    '│                      用户浏览器 (React 19)                        │',
    '│   聊天 │ 资源 │ 路径 │ 辅导 │ 评估 │ 课堂/PPT                   │',
    '│                      api.ts (REST + SSE)                        │',
    '└───────────────────────────────┬──────────────────────────────────┘',
    '                           │ HTTP / SSE / WebSocket',
    '┌───────────────────────────────┴──────────────────────────────────┐',
    '│                    FastAPI 后端 (Python)                         │',
    '│   API 路由层 (20 个路由模块) → 服务层 (16 个服务模块)            │',
    '│   多智能体编排引擎 │ 知识图谱引擎 │ 画像引擎 │ 概念本体引擎     │',
    '│   数据持久层 (SQLAlchemy + SQLite + ChromaDB + 文件系统)         │',
    '└─────────────────────────────────────────────────────────────────────┘',
    '              │              │              │',
    '   ┌──────────┴─┐  ┌─────────┴───┐  ┌─────┴──────┐',
    '   │ 科大讯飞    │  │ DuckDuckGo   │  │ 讯飞 ChatDoc │',
    '   │ 星火大模型  │  │ + Bing       │  │ 知识库 API   │',
    '   └──────────────┘  └────────────────┘  └──────────────┘',
]
add_code_block(arch)

add_heading('3.2 数据流设计', level=2)

add_heading('3.2.1 对话画像流（Chat → Profile）', level=3)
for item in [
    '1. 用户输入消息 → FastAPI 接收 (chat.py)',
    '2. 加载对话历史 + 学生画像 (DB)',
    '3. RAG 检索相关知识 (ChromaDB + Spark KB)',
    '4. 实时学习状态分析 (realtime_state_service)',
    '5. 构建系统提示词 (prompts.py → chat_system())',
    '6. LLM 流式生成 (spark_service.chat_stream())',
    '7. SSE 推送文本到前端',
    '8. ProfileCoordinator 提取画像 → 写入 DB → SSE 推送',
]:
    add_para(item)

add_heading('3.2.2 资源生成流（Multi-Agent Pipeline）', level=3)
for item in [
    '1. 用户输入主题 + 选择类型',
    '2. ResourceCoordinator 编排 4 步 Pipeline:',
    '   Step 1: RAG 检索 (课程知识库 + 用户文档 + Spark KB + 联网搜索)',
    '   Step 2: Orchestrator Agent 需求分析 + 资源规划',
    '   Step 3: 专业 Agent 内容生成 (doc/mindmap/quiz/video/code)',
    '   Step 4: Safety Agent 双层安全审查',
    '3. 每步通过 SSE 推送 agent_status 事件',
    '4. 最终内容通过 SSE text 事件流式输出',
]:
    add_para(item)

add_heading('3.2.3 学习路径生成流', level=3)
for item in [
    '1. 用户请求生成路径',
    '2. 加载学生画像 + 知识图谱 (10 章 DAG)',
    '3. LLM 生成个性化路径结构',
    '4. Kahn 拓扑排序 + 画像权重调整',
    '5. SSE 推送路径数据 + 分析文本',
    '6. 前端渲染: 时间线视图 / 交互式知识图谱 (SVG)',
]:
    add_para(item)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 4
# ══════════════════════════════════════════════════
add_heading('四、技术架构', level=1)

add_heading('4.1 技术栈总览', level=2)
add_table(
    ['层级', '技术选型', '版本', '选型理由'],
    [
        ['前端框架', 'React', '19.2.6', '组件化开发，生态成熟'],
        ['类型系统', 'TypeScript', '6.x', '类型安全，提升代码质量'],
        ['构建工具', 'Vite', '8.0.12', '极速 HMR，原生 ESM'],
        ['样式框架', 'Tailwind CSS', '4.3.0', '原子化 CSS，快速迭代'],
        ['状态管理', 'Zustand', '5.0.14', '轻量级，TypeScript 友好'],
        ['路由', 'React Router', '7.17.0', '声明式路由，嵌套布局'],
        ['后端框架', 'FastAPI', '0.115.0', '异步高性能，自动文档'],
        ['ORM', 'SQLAlchemy', '2.0.35', '异步支持，类型安全'],
        ['数据库', 'SQLite + aiosqlite', '-', '轻量部署，异步读写'],
        ['向量库', 'ChromaDB', '0.5.23', '嵌入式向量检索'],
        ['LLM', '科大讯飞星火大模型', 'spark-x', '对话/生成/辅导'],
        ['Embedding', '星火 Embedding API', 'embedding-2', '1024 维向量化'],
        ['认证', 'JWT', '-', 'python-jose + bcrypt'],
    ],
    [2, 3.5, 2.5, 6]
)

add_heading('4.2 前端技术栈详解', level=2)
add_table(
    ['技术', '版本', '用途'],
    [
        ['react-markdown', '10.1.0', 'Markdown 渲染'],
        ['remark-gfm', '4.0.1', 'GitHub 风格 Markdown（表格、任务列表）'],
        ['remark-math', '6.0.0', '数学公式解析'],
        ['rehype-katex', '7.0.1', 'KaTeX 数学公式渲染'],
        ['KaTeX', '0.17.0', '数学公式渲染引擎'],
        ['Mermaid', '11.15.0', '图表渲染（流程图、思维导图）'],
        ['Prism.js', '1.30.0', '代码语法高亮'],
        ['rehype-prism-plus', '2.0.2', '代码高亮增强插件'],
    ],
    [3.5, 2, 8.5]
)

add_heading('4.3 后端技术栈详解', level=2)
add_table(
    ['技术', '版本', '用途'],
    [
        ['Uvicorn', '0.30.6', 'ASGI 服务器'],
        ['httpx', '0.27.2', '异步 HTTP 客户端'],
        ['websockets', '≥13.0', 'WebSocket 客户端（讯飞 API）'],
        ['python-pptx', '≥1.0.0', '本地 PPT 生成'],
        ['pypdf', '≥4.0.0', 'PDF 文档解析'],
        ['python-docx', '≥1.1.0', 'Word 文档解析'],
        ['jieba', '0.42.1', '中文分词（安全过滤）'],
        ['CrewAI', '0.80.0', '多智能体框架定义'],
        ['pydantic', '2.9.2', '数据校验与序列化'],
    ],
    [3.5, 2, 8.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 5
# ══════════════════════════════════════════════════
add_heading('五、数据模型设计', level=1)

add_heading('5.1 数据库表结构概览', level=2)
add_para('系统共设计 14 张数据表，使用 SQLAlchemy ORM 定义，SQLite 异步引擎驱动。核心表包括 User、StudentProfile、Conversation、ConversationMessage、LearningResource、LearningPath、AssessmentRecord，以及 ConceptMastery、RealtimeLearningState、Flashcard、UserDocument、SavedCourse、FeynmanRecord、PPTRecord 等扩展表。')

add_heading('5.2 核心数据表详解', level=2)

add_heading('5.2.1 User（用户账号表）', level=3)
add_table(
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'Integer', 'PK, autoincrement', '主键'],
        ['username', 'String(64)', 'UNIQUE, NOT NULL', '用户名'],
        ['hashed_password', 'String(256)', 'NOT NULL', 'bcrypt 密码哈希'],
        ['nickname', 'String(64)', '-', '显示昵称'],
        ['created_at', 'DateTime', 'default=now', '创建时间'],
    ],
    [3, 2.5, 3.5, 5]
)

add_heading('5.2.2 StudentProfile（六维学习画像表）', level=3)
add_table(
    ['字段', '类型', '说明'],
    [
        ['id', 'Integer', '主键'],
        ['user_id', 'String(64)', '用户标识，唯一索引'],
        ['knowledge_base', 'JSON', '知识基础：{子领域: 掌握度(0-1)}'],
        ['cognitive_style', 'String(32)', '认知风格：visual/verbal/active/reflective'],
        ['weak_points', 'JSON', '易错点列表'],
        ['learning_goal', 'Text', '学习目标描述'],
        ['available_time', 'String(32)', '可用学习时间'],
        ['interests', 'JSON', '兴趣方向列表'],
        ['conversation_summary', 'Text', '对话历史摘要'],
    ],
    [3, 2.5, 8.5]
)

add_para('六维画像维度详解：', bold=True)
add_table(
    ['维度', '字段', '数据格式', '提取方式'],
    [
        ['知识基础', 'knowledge_base', '{\"机器学习\": 0.6, ...}', 'LLM 从对话推断'],
        ['认知风格', 'cognitive_style', '\"visual\" / \"verbal\" 等', 'LLM 从对话推断'],
        ['易错点', 'weak_points', '[\"梯度消失\", \"过拟合\"]', '做题诊断 + 对话分析'],
        ['学习目标', 'learning_goal', '\"掌握深度学习基础\"', '对话显式获取'],
        ['可用时间', 'available_time', '\"每天 2 小时\"', '对话显式获取'],
        ['兴趣方向', 'interests', '[\"计算机视觉\", \"NLP\"]', 'LLM 从对话推断'],
    ],
    [2, 2.5, 5, 4.5]
)

add_heading('5.2.3 其他数据表', level=3)
add_table(
    ['表名', '说明'],
    [
        ['Conversation', '对话会话表：user_id + title + created_at'],
        ['ConversationMessage', '对话消息表：conversation_id + role + content'],
        ['LearningResource', '学习资源表：resource_type + title + content + chapter + difficulty'],
        ['LearningPath', '学习路径表：path_data(JSON) + current_node + completed_nodes + progress'],
        ['AssessmentRecord', '学习评估记录表：study_time + quiz_scores + assessment_report'],
        ['ConceptMastery', '概念掌握度追踪：含 SM-2 参数 (ease_factor, interval_days)'],
        ['RealtimeLearningState', '实时学习状态：emotion, confusion, cognitive_load, confidence, engagement'],
        ['Flashcard', '闪卡复习：SM-2 算法调度参数'],
        ['UserDocument', '用户知识库文档：多格式、标签、笔记、Spark KB 同步'],
        ['SavedCourse', 'AI 课堂课程：大纲 + 已完成课时追踪'],
        ['FeynmanRecord', '费曼学习法对话记录：含理解度评分'],
        ['PPTRecord', 'PPT 生成历史：讯飞 CDN URL 或本地路径'],
    ],
    [3.5, 10.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 6
# ══════════════════════════════════════════════════
add_heading('六、多智能体系统设计', level=1)

add_heading('6.1 智能体架构总览', level=2)
add_para('系统设计了 13 个 AI Agent，分布在 4 个协作链路中：')
add_table(
    ['协作链路', '编排器', 'Agent 组成'],
    [
        ['资源生成', 'ResourceCoordinator', 'RAG 检索助手 → Orchestrator → 专业 Agent → 安全审查'],
        ['画像提取', 'ProfileCoordinator', 'ProfileExtractor → 画像写入 DB → SSE 推送更新'],
        ['学习路径', 'PathPlanner', 'LLM 结构生成 → 知识图谱拓扑排序 → 个性化权重调整'],
        ['智能辅导', 'TutorAgent', 'RAG 多源检索 → 自适应难度生成 → 来源引用'],
    ],
    [2.5, 3.5, 8]
)

add_heading('6.2 资源生成编排器（ResourceCoordinator）', level=2)
add_para('核心文件：agents/coordinator.py（787 行）。ResourceCoordinator 是资源生成的核心编排器，实现了 4 步 Pipeline 模式。')

add_heading('Step 1：RAG 多源检索', level=3)
add_para('四路并行检索并合并去重：课程知识库 (ChromaDB)、用户文档库、Spark 知识库 (讯飞 ChatDoc)、联网搜索 (DuckDuckGo + Bing)。')

add_heading('Step 2：Orchestrator Agent 需求分析', level=3)
add_para('角色为资源设计总监（Resource Design Director），分析用户需求，制定资源生成策略。')

add_heading('Step 3：专业 Agent 内容生成', level=3)
add_table(
    ['资源类型', 'Agent 名称', '角色定义', '生成内容'],
    [
        ['doc', '课程内容专家', 'Course Content Expert', 'Markdown 讲义（含 LaTeX、Mermaid）'],
        ['mindmap', '知识建筑师', 'Knowledge Architect', 'Mermaid 思维导图'],
        ['quiz', '题目设计专家', 'Quiz Design Expert', '选择/判断/简答/编程题'],
        ['video', '多媒体编剧', 'Multimedia Scriptwriter', '视频脚本（分镜、旁白、时长）'],
        ['code', '编程实践导师', 'Programming Practice Mentor', 'Python 编程练习 + 测试用例'],
    ],
    [2, 2.5, 4, 5.5]
)

add_heading('Step 4：Safety Agent 双层安全审查', level=3)
add_para('第一层：正则表达式快筛（10 类敏感模式，< 10ms）。第二层：LLM-as-Judge 深度审查（内容安全性 + 幻觉检测）。')

add_heading('6.3 画像提取编排器（ProfileCoordinator）', level=2)
add_para('核心文件：agents/profile_coordinator.py（93 行）。提取流程：取最近 6 条对话消息 → 格式化对话文本 → 包含现有画像上下文 → LLM 提取 JSON（temperature=0.1）→ 鲁棒性 JSON 解析 → 返回 7 字段画像字典。')

add_para('动态节流策略：', bold=True)
add_table(
    ['画像完整度', '提取频率', '理由'],
    [
        ['不完整（< 3 维有值）', '每轮对话提取', '尽快构建完整画像'],
        ['基本完整（3-5 维）', '每 3 轮提取一次', '平衡精度与效率'],
        ['完整（6 维全有值）', '每 5 轮提取一次', '避免不必要的 LLM 调用'],
    ],
    [4, 3.5, 6.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 7
# ══════════════════════════════════════════════════
add_heading('七、前端设计与实现', level=1)

add_heading('7.1 页面结构', level=2)
add_para('系统共 16 个页面，对应 14 个侧边栏导航项：')
add_table(
    ['序号', '页面', '路由', '核心功能'],
    [
        ['1', 'Dashboard', '/', '统计概览、知识掌握度、快捷操作、成就徽章'],
        ['2', '对话学习', '/chat', 'SSE 流式对话 + 画像自动提取 + 建议回复'],
        ['3', 'AI 课堂', '/classroom', 'AI 生成课件、教师讲稿、课堂测验'],
        ['4', 'AI 辅导', '/tutoring', 'RAG 智能答疑 + 来源引用 + Mermaid 图解'],
        ['5', '闪卡复习', '/flashcards', 'SM-2 间隔重复 + 翻转卡片'],
        ['6', '在线练习', '/quiz', '章节选择 + AI 出题 + 自动批改'],
        ['7', '错题本', '/wrong-answer-book', '错题收录 + 复习模式 + 标记已掌握'],
        ['8', '学习资源', '/resources', '5 类资源生成 + 列表管理 + 搜索收藏'],
        ['9', '知识库', '/knowledge-base', '文档上传/导入 + 向量同步 + AI 问答'],
        ['10', '学习路径', '/learning-path', '时间线 + 交互式知识图谱 + 节点详情'],
        ['11', '学习旅程', '/learning-journey', '保存的课程 + 费曼学习法'],
        ['12', '学习评估', '/assessment', '五维雷达图 + AI 报告 + 趋势折线图'],
        ['13', 'PPT 生成', '/ppt', 'AI 大纲生成 + 一键生成 PPT'],
        ['14', '个人画像', '/profile', '画像查看/编辑 + 历史变化曲线'],
    ],
    [1, 2.5, 3.5, 7]
)

add_heading('7.2 可复用组件（17 个）', level=2)
add_table(
    ['组件', '文件', '功能'],
    [
        ['AppLayout', 'AppLayout.tsx', '根布局包装器'],
        ['Sidebar', 'Sidebar.tsx', '可折叠导航侧边栏（深色渐变 + 星空特效）'],
        ['ChatInput', 'ChatInput.tsx', '聊天输入框'],
        ['ChatBubble', 'ChatBubble.tsx', '聊天气泡（Markdown 渲染）'],
        ['CodeBlock', 'CodeBlock.tsx', '语法高亮代码块（Prism.js）'],
        ['MermaidBlock', 'MermaidBlock.tsx', 'Mermaid 图表渲染器（缩放/拖拽/缓存）'],
        ['ProfilePanel', 'ProfilePanel.tsx', '学生画像侧边面板'],
        ['ErrorBoundary', 'ErrorBoundary.tsx', 'React 错误边界'],
        ['DigitalHumanAvatar', 'DigitalHumanAvatar.tsx', '数字人头像组件'],
        ['VoiceModePanel', 'VoiceModePanel.tsx', '语音交互面板（音频可视化）'],
        ['FloatingProgressPanel', 'FloatingProgressPanel.tsx', '浮动进度面板'],
        ['FeynmanPanel', 'FeynmanPanel.tsx', '费曼学习法交互面板'],
        ['ImageUploader', 'ImageUploader.tsx', '图片上传预览组件'],
    ],
    [3, 4, 7]
)

add_heading('7.3 状态管理（Zustand Store）', level=2)
add_table(
    ['Store', '文件', '职责'],
    [
        ['authStore', 'authStore.ts', '认证状态：token、用户信息、登录/登出'],
        ['profileStore', 'profileStore.ts', '学生画像：按 username 隔离缓存，localStorage 持久化'],
        ['chatStore', 'chatStore.ts', '当前对话消息状态'],
        ['classroomStore', 'classroomStore.ts', '课堂会话状态（幻灯片、讲稿、测验）'],
    ],
    [2.5, 3.5, 8]
)

add_heading('7.4 设计风格', level=2)
add_table(
    ['设计要素', '规范'],
    [
        ['主色调', '深墨绿 (#1A3A2A) + 暖白 (#FAF8F5)'],
        ['强调色', '琥珀色 (#D97706)'],
        ['字体', '系统字体栈，正文 14px，辅助 12px'],
        ['圆角', '卡片 12px，按钮 8px，输入框 6px'],
        ['阴影', '柔和投影 shadow-sm'],
        ['动画', 'fade-in 过渡 (0.3s ease-out)'],
    ],
    [3.5, 10.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 8
# ══════════════════════════════════════════════════
add_heading('八、后端设计与实现', level=1)

add_heading('8.1 核心服务详解', level=2)

add_heading('8.1.1 SparkService（LLM 客户端）', level=3)
add_para('接入讯飞星火大模型 REST API（模型 spark-x），提供同步 chat() 和异步 chat_stream() 两种调用方式，超时 120 秒，支持自动重试。')

add_heading('8.1.2 RAGService（RAG 检索服务）', level=3)
add_para('基于 ChromaDB PersistentClient，管理课程知识库（10 章 Markdown，按 ## 标题分块）和用户文档集合（按用户隔离）。支持语义向量检索和文本模糊匹配兜底，自动检测嵌入维度不匹配并重建集合。')

add_heading('8.1.3 SafetyService（安全审查服务）', level=3)
add_para('第一层正则快筛覆盖 10 类敏感模式（暴力、武器、犯罪、自残、歧视、色情、政治、隐私、医疗、金融），附加幻觉标记检测和外部图片 URL 检测。第二层 LLM-as-Judge 提供内容安全审查和幻觉检测，采用 fail-closed 设计。')

add_heading('8.1.4 EmbeddingService（向量化服务）', level=3)
add_para('主方案为讯飞星火 Embedding API（embedding-2，1024 维），降级方案为本地 n-gram 哈希向量化。自动检测集合嵌入方法不一致时触发重建。')

add_heading('8.2 集中化提示词管理', level=2)
add_para('核心文件：backend/prompts.py（714 行）。所有 LLM 提示词集中管理，分为 14 个功能模块：')
add_table(
    ['模块', '使用场景'],
    [
        ['对话系统', '对话画像构建（CHAT_SYSTEM_BASE 等）'],
        ['辅导系统', '智能辅导答疑（TUTORING_SYSTEM_BASE 等）'],
        ['画像提取', '从对话提取六维画像（EXTRACT_SYSTEM_PROMPT）'],
        ['资源生成', '多智能体资源生成（DOC/MINDMAP/QUIZ/VIDEO/CODE_AGENT_ROLE）'],
        ['学习路径', '路径生成与分析（PATH_STRUCTURE_PROMPT 等）'],
        ['学习评估', '评估报告生成（ASSESSMENT_SYSTEM 等）'],
        ['安全审查', 'LLM 安全审查（SAFETY_JUDGE_PROMPT 等）'],
        ['PPT 生成', 'PPT 大纲生成（OUTLINE_PROMPT 等）'],
        ['错误诊断', '薄弱点诊断（DIAGNOSTIC_SYSTEM 等）'],
        ['闪卡生成', 'SM-2 闪卡生成（FLASHCARD_GENERATION_PROMPT）'],
        ['每日计划', '日计划生成（DAILY_PLAN_SYSTEM 等）'],
        ['费曼学习', '苏格拉底式对话（FEYNMAN_SYSTEM 等）'],
        ['门卫微课', '前置知识补课（GATEKEEPER_MINILESSON_PROMPT）'],
        ['自适应教学', '根据画像生成教学指令 (build_adaptive_instruction())'],
    ],
    [3.5, 10.5]
)
add_para('设计原则：统一人设（“你是智学 [角色名]”）、深度/快速双模式、内部策略标记（“禁止向用户暴露”）。')

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 9
# ══════════════════════════════════════════════════
add_heading('九、API 接口规范', level=1)

add_heading('9.1 认证方式', level=2)
add_para('采用 JWT Bearer Token 认证，Header 格式为 Authorization: Bearer <token>，Token 有效期 24 小时，通过 /api/auth/login 或 /api/auth/register 获取。')

add_heading('9.2 SSE 流式接口规范', level=2)
add_para('所有 AI 生成接口均使用 Server-Sent Events 流式输出，Content-Type 为 text/event-stream。事件类型包括：')
add_table(
    ['事件类型', '说明', '使用场景'],
    [
        ['text', '增量文本内容', '所有流式接口'],
        ['done', '流结束标记', '所有流式接口'],
        ['error', '错误信息', '所有流式接口'],
        ['warning', '安全警告', '资源生成、辅导'],
        ['agent_status', 'Agent 工作状态', '资源生成'],
        ['profile_update', '画像更新数据', '对话'],
        ['path_data', '路径结构数据', '学习路径'],
        ['realtime_state', '实时学习状态', '对话'],
    ],
    [3, 4.5, 6.5]
)

add_heading('9.3 API 端点清单', level=2)

api_modules = [
    ('9.3.1 认证模块', [
        ['POST', '/api/auth/register', '用户注册', '无需'],
        ['POST', '/api/auth/login', '用户登录', '无需'],
        ['GET', '/api/auth/me', '当前用户信息', 'Bearer'],
    ]),
    ('9.3.2 对话模块', [
        ['POST', '/api/chat/stream', 'SSE 流式对话 + 画像提取', 'Bearer'],
        ['GET', '/api/conversations/', '对话列表', 'Bearer'],
        ['POST', '/api/conversations/', '创建对话', 'Bearer'],
        ['GET', '/api/conversations/{id}/messages', '获取对话消息', 'Bearer'],
        ['DELETE', '/api/conversations/{id}', '删除对话', 'Bearer'],
    ]),
    ('9.3.3 学习资源模块', [
        ['POST', '/api/resources/generate', 'SSE 流式多智能体资源生成', 'Bearer'],
        ['GET', '/api/resources/', '资源列表（画像排序）', 'Bearer'],
        ['GET', '/api/resources/{id}', '资源详情', 'Bearer'],
        ['PUT', '/api/resources/{id}', '更新资源', 'Bearer'],
        ['DELETE', '/api/resources/{id}', '删除资源', 'Bearer'],
    ]),
    ('9.3.4 学习路径模块', [
        ['POST', '/api/learning-path/generate', 'SSE 流式路径生成', 'Bearer'],
        ['GET', '/api/learning-path/', '获取路径', 'Bearer'],
        ['POST', '/api/learning-path/toggle-node', '切换节点完成状态', 'Bearer'],
    ]),
]

for title, rows in api_modules:
    add_heading(title, level=3)
    add_table(['方法', '端点', '说明', '认证'], rows, [1.5, 4.5, 5.5, 2.5])

add_heading('9.3.5 其他模块', level=3)
add_table(
    ['模块', '端点数', '主要端点'],
    [
        ['学生画像', '2', 'GET /api/profile/ , POST /api/profile/update'],
        ['学习评估', '4', 'generate, record, trends, get'],
        ['智能辅导', '1', 'POST /api/tutoring/ask (SSE)'],
        ['知识库', '10', 'upload, import-folder, web-search, documents, search, ask'],
        ['PPT 生成', '7', 'outline, create-from-outline, create, progress, download, records'],
        ['AI 课堂', '8', 'outline, start, submit-practice, save, courses, feynman'],
        ['其他', '12', 'mastery, daily-plan, flashcards, learning-journey, diagnose, realtime-state'],
    ],
    [2.5, 2, 9.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 10
# ══════════════════════════════════════════════════
add_heading('十、知识库与 RAG 系统', level=1)

add_heading('10.1 课程知识库', level=2)
add_para('以“人工智能导论”为内容，构建 10 章完整知识体系：')
add_table(
    ['章节', '主题', '难度', '核心内容'],
    [
        ['ch01', '人工智能导论', '0.20', 'AI 定义、发展历程、学派、应用领域'],
        ['ch02', '机器学习基础', '0.45', '监督/无监督/强化学习、回归、决策树'],
        ['ch03', '深度学习基础', '0.55', '神经网络、反向传播、激活函数'],
        ['ch04', 'Transformer 架构', '0.65', '注意力机制、BERT/GPT'],
        ['ch05', '自然语言处理', '0.60', '分词、词嵌入、语言模型'],
        ['ch06', '计算机视觉', '0.60', 'CNN、目标检测、GAN'],
        ['ch07', '强化学习', '0.70', 'MDP、Q-Learning、策略梯度'],
        ['ch08', 'AI 伦理与安全', '0.30', '偏见、公平性、隐私、法规'],
        ['ch09', 'MLOps', '0.50', '部署、监控、CI/CD'],
        ['ch10', '前沿方向', '0.50', '多模态 AI、基础模型'],
    ],
    [1.5, 3, 1.5, 8]
)

add_heading('10.2 知识图谱引擎', level=2)
add_para('核心文件：knowledge_base/knowledge_graph.py（638 行）。10 章 DAG 依赖关系：ch01 (AI 导论) → ch02 (ML) → ch03 (DL) → ch04 (Transformer) → ch05 (NLP) / ch06 (CV)；ch03 → ch07 (RL)；ch02 → ch09 (MLOps)；ch01 → ch08 (伦理) / ch10 (前沿)。')
add_para('使用 Kahn 算法进行拓扑排序确保前置依赖，再根据学生画像的掌握度、易错点、兴趣方向和学习目标进行个性化权重调整。')

add_heading('10.3 向量检索流程', level=2)
add_para('用户查询 → 文本预处理（jieba 分词）→ Embedding 向量化（星火 API 1024 维）→ ChromaDB 余弦相似度检索 → Top-K 结果返回 → 注入 LLM 系统提示词作为上下文。')

add_heading('10.4 用户知识库管理', level=2)
add_para('支持 6 种文档格式：.md（直接读取）、.txt（直接读取）、.pdf（pypdf 解析）、.docx（python-docx 解析）、.html（直接读取）、.csv（直接读取）。文档处理流程：上传 → 格式解析 → 分块（~500 字/块）→ 向量化 → 存入 ChromaDB → 同步到 Spark ChatDoc（可选）。')

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 11
# ══════════════════════════════════════════════════
add_heading('十一、安全设计', level=1)

add_heading('11.1 认证安全', level=2)
add_table(
    ['安全措施', '实现方式'],
    [
        ['密码存储', 'bcrypt 单向哈希'],
        ['Token 签发', 'JWT HS256，24 小时有效期'],
        ['密钥管理', '.env 配置，不入 Git；启动时 fail-closed 校验（拒绝默认密钥启动）'],
        ['输入校验', 'Pydantic 校验（用户名 2-20 位字母数字，密码 4-20 位）'],
        ['前端安全', '无 console.log（防密码泄露），401 自动跳转登录'],
    ],
    [3.5, 10.5]
)

add_heading('11.2 内容安全', level=2)
add_para('双层安全审查架构：第一层正则快筛（< 10ms，10 类敏感模式 + 幻觉标记 + 外部图片 URL），通过后进入第二层 LLM-as-Judge（内容安全性审查 + 幻觉检测，fail-closed）。通过后添加 AI 生成声明。')

add_heading('11.3 数据安全', level=2)
add_table(
    ['安全措施', '说明'],
    [
        ['用户数据隔离', '所有数据按 user_id 隔离查询'],
        ['资源所有权校验', '资源操作需验证所有权'],
        ['.env gitignore', '敏感配置不入版本控制'],
        ['CORS 白名单', '仅允许指定域名访问'],
    ],
    [3.5, 10.5]
)

add_heading('11.4 LLM 安全防护', level=2)
add_table(
    ['防护措施', '说明'],
    [
        ['System Prompt 注入防护', '提示词模板化，用户输入与指令隔离'],
        ['幻觉防范', 'RAG 检索提供事实依据 + 幻觉检测审查'],
        ['AI 生成声明', '所有生成内容末尾添加“以上内容由 AI 生成，仅供参考”'],
        ['内容过滤', '10 类正则模式 + LLM 审查双层过滤'],
    ],
    [3.5, 10.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 12
# ══════════════════════════════════════════════════
add_heading('十二、核心算法与服务', level=1)

add_heading('12.1 SM-2 间隔重复算法', level=2)
add_para('用于闪卡复习调度，基于遗忘曲线原理。核心参数：ease_factor（易度因子，≥ 1.3）、interval_days（间隔天数）、next_review_at（下次复习时间）。回忆质量 quality ≥ 3 时增大间隔，quality < 3 时重置为 1 天。')

add_heading('12.2 概念本体系统', level=2)
add_para('核心文件：backend/concept_ontology.py（500 行）。定义 30+ 个 AI 核心概念的形式化知识体系，每个概念包含 concept_id、name、chapter、difficulty、prerequisites（前置概念）、related_concepts（关联概念）、keywords（关键词）。应用场景包括学习路径依赖计算、薄弱点诊断、推荐相关知识点。')

add_heading('12.3 实时学习状态分析', level=2)
add_para('核心文件：backend/services/realtime_state_service.py（236 行）。双层分析策略：第一层规则匹配（基于关键词的情绪/困惑检测，无 LLM 调用，< 10ms），第二层 LLM 增强（复杂语义分析，按需触发）。追踪 5 维实时状态：emotion（positive/neutral/negative）、confusion（0-1）、cognitive_load（low/medium/high）、confidence（0-1）、engagement（low/medium/high）。')

add_heading('12.4 薄弱点诊断', level=2)
add_para('收集做题记录 → 错误率 > 50% 标记为薄弱点 → 概念本体检查前置概念 → 自动更新 weak_points → LLM 生成诊断报告（错误类型分类：概念混淆 / 知识空白 / 粗心错误 / 推理偏差）。')

add_heading('12.5 每日学习计划', level=2)
add_para('生成考虑因素：学生画像（知识基础、学习目标、可用时间）、当前学习路径进度、待复习闪卡数量、薄弱点优先级、概念本体前置依赖。')

add_heading('12.6 概念掌握度追踪', level=2)
add_para('基于遗忘曲线的掌握度衰减模型：mastery(t) = mastery_0 × e^(-λ × t)，其中 λ 为 forgetting_factor（因人而异，根据复习频率调整）。每次复习后 mastery 更新并重置衰减起点。')

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 13
# ══════════════════════════════════════════════════
add_heading('十三、部署与运维', level=1)

add_heading('13.1 环境要求', level=2)
add_table(
    ['组件', '版本要求', '说明'],
    [
        ['Python', '≥ 3.11', '后端运行环境'],
        ['Node.js', '≥ 20', '前端构建环境'],
        ['操作系统', 'Windows / macOS / Linux', '跨平台支持'],
    ],
    [3.5, 3.5, 7]
)

add_heading('13.2 配置步骤', level=2)
steps = [
    ('步骤 1：克隆项目', 'git clone <repository-url> && cd chinese.software.cup2'),
    ('步骤 2：配置环境变量', 'cd backend && cp .env.example .env，编辑 .env 填写讯飞星火 API 凭证。必填项：SPARK_APP_ID、SPARK_API_KEY、SPARK_API_SECRET、JWT_SECRET_KEY。'),
    ('步骤 3：启动后端', 'cd backend && pip install -r requirements.txt && python run.py（后端启动于 http://localhost:8000）。启动时自动创建数据库表、执行 SQLite 迁移、校验 JWT 密钥、后台构建课程知识库。'),
    ('步骤 4：启动前端', 'cd frontend && npm install && npm run dev（前端启动于 http://localhost:5173）。'),
]
for title, desc in steps:
    para = doc.add_paragraph()
    r1 = para.add_run(title + '\n')
    r1.bold = True
    r2 = para.add_run(desc)

add_heading('13.3 知识库重建', level=2)
add_para('升级后如需重建向量索引，执行 python -c "from knowledge_base.build_kb import build_knowledge_base; import asyncio; asyncio.run(build_knowledge_base())"。')

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 14
# ══════════════════════════════════════════════════
add_heading('十四、测试方案', level=1)

add_heading('14.1 测试环境', level=2)
add_table(
    ['项目', '配置'],
    [
        ['操作系统', 'Windows 11'],
        ['Python', '3.11+'],
        ['Node.js', '22.x'],
        ['浏览器', 'Chrome / Edge（最新版）'],
        ['LLM', '讯飞星火 spark-x'],
    ],
    [3.5, 10.5]
)

add_heading('14.2 功能测试矩阵', level=2)

test_groups = [
    ('FC-01：对话式画像构建（6 项）', [
        ['FC-01-01', '新用户首条消息', '自动创建对话，AI 流式回复'],
        ['FC-01-02', '多轮对话', '画像自动更新（SSE profile_update 事件）'],
        ['FC-01-03', '对话列表', '正确显示历史对话，支持删除'],
        ['FC-01-04', '切换对话', '正确加载历史消息'],
        ['FC-01-05', '建议回复按钮', 'AI 回复中提取建议，显示可点击按钮'],
        ['FC-01-06', '画像手动编辑', 'Profile 页面修改后保存到 DB'],
    ]),
    ('FC-02：多智能体资源生成（7 项）', [
        ['FC-02-01', '生成文档资源', '4 个 Agent 依次显示状态，最终输出 Markdown'],
        ['FC-02-02', '生成思维导图', '输出 Mermaid 格式，前端正确渲染'],
        ['FC-02-03', '生成练习题', '输出选择/判断/简答题，含答案解析'],
        ['FC-02-04', '生成视频脚本', '输出分镜脚本格式'],
        ['FC-02-05', '生成代码示例', '输出 Python 代码 + 测试用例'],
        ['FC-02-06', 'Agent 状态面板', '4 个 Agent 图标依次 working → done'],
        ['FC-02-07', '安全审查', '含违规内容的请求被拦截'],
    ]),
    ('FC-03：学习路径规划（3 项）', [
        ['FC-03-01', '生成路径', '显示时间线视图 + 知识图谱'],
        ['FC-03-02', '节点完成', '点击标记已完成，进度更新'],
        ['FC-03-03', '知识图谱交互', 'hover 高亮、缩放、拖拽、节点详情面板'],
    ]),
    ('FC-04：智能辅导（3 项）', [
        ['FC-04-01', '提问答疑', '流式回答 + 来源引用 [1][2]'],
        ['FC-04-02', '图表生成', '回答中包含 Mermaid 图表'],
        ['FC-04-03', '追问建议', '回答后显示 2-3 个追问按钮'],
    ]),
    ('FC-05：学习评估（4 项）', [
        ['FC-05-01', '生成评估报告', '五维雷达图 + 文字报告'],
        ['FC-05-02', '学习趋势图', '14 天折线图'],
        ['FC-05-03', '知识点热力图', '各章节掌握度可视化'],
        ['FC-05-04', '报告导出', 'PDF 打印导出'],
    ]),
]

for title, rows in test_groups:
    add_heading(title, level=3)
    add_table(['编号', '测试项', '预期结果'], rows, [2, 3.5, 8.5])

add_heading('14.3 非功能测试', level=2)
add_table(
    ['测试类型', '测试内容', '通过标准'],
    [
        ['流式延迟', 'SSE 首 token 到达时间', '< 2 秒'],
        ['Markdown 渲染', '表格/列表/代码块', '正确渲染'],
        ['LaTeX 渲染', '数学公式', 'KaTeX 正确显示'],
        ['Mermaid 渲染', '流程图/思维导图', 'SVG 正确生成'],
        ['空状态处理', '无数据时的界面', '显示引导文案'],
        ['错误处理', 'API 异常', '显示友好错误提示'],
        ['内容安全', '违规输入', '被安全过滤拦截'],
        ['前端构建', 'npm run build', '无错误'],
    ],
    [3, 5, 6]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 15
# ══════════════════════════════════════════════════
add_heading('十五、开源工具使用标注', level=1)

add_heading('15.1 前端依赖', level=2)
add_table(
    ['工具', '版本', '协议', '用途'],
    [
        ['React', '19.2.6', 'MIT', 'UI 框架'],
        ['React Router', '7.17.0', 'MIT', '路由管理'],
        ['Zustand', '5.0.14', 'MIT', '状态管理'],
        ['Tailwind CSS', '4.3.0', 'MIT', '样式框架'],
        ['Vite', '8.0.12', 'MIT', '构建工具'],
        ['TypeScript', '6.x', 'Apache-2.0', '类型系统'],
        ['react-markdown', '10.1.0', 'MIT', 'Markdown 渲染'],
        ['remark-gfm', '4.0.1', 'MIT', 'GFM 扩展'],
        ['remark-math', '6.0.0', 'MIT', '数学公式解析'],
        ['rehype-katex', '7.0.1', 'MIT', 'KaTeX 渲染'],
        ['KaTeX', '0.17.0', 'MIT', '数学公式引擎'],
        ['Mermaid', '11.15.0', 'MIT', '图表渲染'],
        ['Prism.js', '1.30.0', 'MIT', '代码高亮'],
        ['rehype-prism-plus', '2.0.2', 'MIT', '代码高亮插件'],
    ],
    [3, 2, 2.5, 6.5]
)

add_heading('15.2 后端依赖', level=2)
add_table(
    ['工具', '版本', '协议', '用途'],
    [
        ['FastAPI', '0.115.0', 'MIT', 'Web 框架'],
        ['Uvicorn', '0.30.6', 'BSD-3', 'ASGI 服务器'],
        ['SQLAlchemy', '2.0.35', 'MIT', 'ORM'],
        ['aiosqlite', '0.20.0', 'Apache-2.0', '异步 SQLite'],
        ['ChromaDB', '0.5.23', 'Apache-2.0', '向量数据库'],
        ['python-jose', '3.3.0', 'MIT', 'JWT 认证'],
        ['passlib', '1.7.4', 'BSD', '密码哈希'],
        ['Pydantic', '2.9.2', 'MIT', '数据校验'],
        ['httpx', '0.27.2', 'BSD-3', 'HTTP 客户端'],
        ['jieba', '0.42.1', 'MIT', '中文分词'],
        ['CrewAI', '0.80.0', 'MIT', '多智能体框架'],
        ['python-pptx', '≥1.0.0', 'MIT', 'PPT 生成'],
        ['pypdf', '≥4.0.0', 'BSD-3', 'PDF 解析'],
        ['python-docx', '≥1.1.0', 'MIT', 'Word 解析'],
    ],
    [3, 2, 2.5, 6.5]
)

add_heading('15.3 AI 服务', level=2)
add_table(
    ['服务', '用途', '说明'],
    [
        ['科大讯飞星火大模型', '主 LLM', '对话/生成/辅导/评估'],
        ['星火 Embedding API', '文本向量化', '1024 维向量'],
        ['星火 ChatDoc API', '云端知识库', '文档问答'],
        ['星火 PPT 生成 API', 'PPT 制作', 'HMAC-SHA256 签名'],
        ['星火数字人 API', '数字形象', 'WebSocket 实时驱动'],
        ['星火 IAT API', '语音识别', '实时语音转文字'],
        ['DuckDuckGo', '联网搜索', '资源生成/辅导补充'],
        ['Bing', '联网搜索', '资源生成/辅导补充'],
    ],
    [3.5, 3, 7.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# CHAPTER 16
# ══════════════════════════════════════════════════
add_heading('十六、项目总结与展望', level=1)

add_heading('16.1 已实现功能总结', level=2)
add_table(
    ['功能模块', '状态', '核心亮点'],
    [
        ['对话式画像构建', '已实现', '六维画像 + 动态节流 + SSE 推送'],
        ['多智能体资源生成', '已实现', '4 Agent 协作 + 5 类资源 + Agent 状态面板'],
        ['个性化学习路径', '已实现', 'Kahn 拓扑排序 + SVG 交互图谱 + 节点详情'],
        ['智能辅导', '已实现', 'RAG + 来源引用 + 难度自适应 + Mermaid 图解'],
        ['学习评估', '已实现', '雷达图 + AI 报告 + 趋势图 + 热力图'],
        ['在线练习', '已实现', '章节选择 + AI 出题 + 自动批改'],
        ['错题本', '已实现', '自动收录 + 复习模式 + SM-2 调度'],
        ['知识库管理', '已实现', '多格式上传 + 向量同步 + AI 问答'],
        ['AI 课堂', '已实现', '课件 + 讲稿 + 测验 + 费曼学习法'],
        ['PPT 生成', '已实现', 'AI 大纲 + 两阶段生成 + 历史记录'],
        ['数字人交互', '已实现', '讯飞数字人 + WebSocket 实时驱动'],
        ['语音交互', '已实现', '实时语音识别 + 语音输入'],
        ['首次引导', '已实现', '4 步向导 + 自动画像构建'],
        ['每日打卡', '已实现', '连续天数 + 学习统计 + 成就徽章'],
        ['闪卡复习', '已实现', 'SM-2 算法 + 翻转卡片 + 间隔重复'],
        ['暗黑模式', '已实现', '深色主题 + CSS 变量切换'],
        ['响应式适配', '已实现', '移动端汉堡菜单 + 自适应布局'],
    ],
    [3.5, 1.5, 9]
)

add_heading('16.2 技术创新点', level=2)
innovations = [
    ('1. 多智能体协同编排', 'ResourceCoordinator 实现 4 步 Pipeline（RAG→Orchestrator→Agent→Safety），支持单阶段和两阶段生成模式'),
    ('2. 六维动态画像', '基于 LLM 的对话式画像提取，配合动态节流策略和增量更新'),
    ('3. 知识图谱 + 画像权重', 'Kahn 拓扑排序确保前置依赖，画像权重实现个性化排序'),
    ('4. 双层安全审查', '正则快筛（< 10ms）+ LLM-as-Judge 深度审查，fail-closed 设计'),
    ('5. 概念本体驱动', '30+ AI 概念的形式化知识体系，支撑路径规划和薄弱点诊断'),
    ('6. 实时学习状态分析', '5 维实时状态追踪（情绪、困惑、认知负荷、自信、参与度）'),
    ('7. SM-2 间隔重复', '基于遗忘曲线的智能复习调度'),
]
for bold_part, desc in innovations:
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    r1 = para.add_run(bold_part)
    r1.bold = True
    para.add_run('：' + desc)

add_heading('16.3 后续优化方向', level=2)
add_table(
    ['方向', '说明'],
    [
        ['Director Graph 编排', '引入 LangGraph 动态决策编排模式'],
        ['多 Agent 讨论', '多角色轮流发言，头像/颜色区分'],
        ['自动导学模式', 'PlaybackEngine 状态机驱动学习路径自动播放'],
        ['AI 白板绘图', '辅导中实时 SVG 白板绘图'],
        ['交互式模拟', 'AI 生成可交互 HTML 组件（神经网络可视化等）'],
    ],
    [3.5, 10.5]
)

# Footer
doc.add_paragraph()
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
pBdr2 = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="8" w:color="1A3A2A"/></w:pBdr>')
para._p.get_or_add_pPr().append(pBdr2)
run = para.add_run('文档编制完成 — 智学（ZhiXue）AI 个性化学习系统 — 2026 年 6 月')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True

# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'docs', '开发说明书.docx')
output_path = os.path.normpath(output_path)
doc.save(output_path)
print(f'Generated: {output_path} ({os.path.getsize(output_path) // 1024} KB)')
