"""文档格式解析服务 — 从 PDF/DOCX/HTML/MD/TXT 中提取纯文本"""

import os
import re
import io
import logging

logger = logging.getLogger(__name__)


def parse_document(content: bytes, filename: str) -> str:
    """根据文件扩展名解析文档内容为纯文本。

    支持: .md .txt .csv .pdf .docx .html
    """
    ext = os.path.splitext(filename or "")[1].lower()

    if ext in (".md", ".txt", ".csv"):
        return _decode_text(content)
    elif ext == ".pdf":
        return _parse_pdf(content)
    elif ext == ".docx":
        return _parse_docx(content)
    elif ext == ".html":
        return _parse_html(content)
    else:
        return _decode_text(content)


def _decode_text(content: bytes) -> str:
    """智能编码检测：chardet 检测 → UTF-8 → GBK → GB18030 → lossy UTF-8"""
    # 先检查是否是 PDF/DOCX 的二进制文件被误当文本传入
    if content[:4] == b'%PDF' or content[:4] == b'PK\x03\x04':
        ext_guess = 'PDF' if content[:4] == b'%PDF' else 'DOCX'
        logger.warning("检测到 %s 二进制文件被传入文本解析器", ext_guess)
        return f"[错误：这是 {ext_guess} 文件，请通过文件上传接口导入]"

    # chardet 检测编码
    try:
        import chardet
        detected = chardet.detect(content)
        if detected and detected.get('encoding') and detected.get('confidence', 0) > 0.5:
            enc = detected['encoding']
            try:
                text = content.decode(enc)
                # 验证解码结果质量（中文内容不应含大量替换字符）
                if '�' not in text[:2000]:
                    return text
            except (UnicodeDecodeError, LookupError):
                pass
    except ImportError:
        pass

    # 逐级降级尝试
    for enc in ('utf-8', 'gbk', 'gb18030', 'big5', 'euc-kr'):
        try:
            text = content.decode(enc)
            if '�' not in text[:2000]:
                return text
        except (UnicodeDecodeError, LookupError):
            continue

    # 最终 lossy 降级
    return content.decode('utf-8', errors='replace')


def _parse_pdf(content: bytes) -> str:
    """从 PDF 提取文本（pypdf 优先，fitz 兜底）"""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        result = "\n\n".join(pages)
        if result.strip():
            return result
        logger.warning("pypdf 未提取到文本，尝试 pymupdf")
    except ImportError:
        logger.info("pypdf 未安装，尝试 pymupdf")
    except Exception as e:
        logger.warning("pypdf 解析失败: %s，尝试 pymupdf", e)

    # pymupdf 兜底
    try:
        import fitz  # type: ignore

        doc = fitz.open(stream=content, filetype="pdf")
        pages = []
        for page in doc:
            text = page.get_text()
            if text:
                pages.append(text.strip())
        doc.close()
        result = "\n\n".join(pages)
        if result.strip():
            return result
    except ImportError:
        logger.error("无 PDF 解析库，请安装: pip install pypdf")
        return "[PDF 解析不可用 — 请安装 pypdf: pip install pypdf]"
    except Exception as e:
        logger.error("pymupdf 解析失败: %s", e)

    return "[PDF 解析失败]"


def _parse_docx(content: bytes) -> str:
    """从 DOCX 提取文本（python-docx）"""
    try:
        from docx import Document  # type: ignore

        doc = Document(io.BytesIO(content))
        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        result = "\n\n".join(paragraphs)
        if result.strip():
            return result
        return "[DOCX 文件内容为空]"
    except ImportError:
        logger.error("python-docx 未安装，请安装: pip install python-docx")
        return "[DOCX 解析不可用 — 请安装 python-docx: pip install python-docx]"
    except Exception as e:
        logger.error("DOCX 解析失败: %s", e)
        return f"[DOCX 解析失败: {e}]"


def _parse_html(content: bytes) -> str:
    """从 HTML 中提取纯文本（去标签）"""
    text = _decode_text(content)
    # 去掉 script/style/nav/footer/header 标签及其内容
    text = re.sub(
        r"<(script|style|nav|footer|header)[^>]*>.*?</\1>",
        "",
        text,
        flags=re.DOTALL | re.IGNORECASE,
    )
    # 去掉 HTML 注释
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    # 去掉所有标签
    text = re.sub(r"<[^>]+>", " ", text)
    # 合并空白
    text = re.sub(r"\s+", " ", text).strip()
    return text
